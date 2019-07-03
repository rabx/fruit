####################################################
__author__	= "rabx"
__version__	= '2.0'
__license__	= 'GNU General Public License v3.0 or later (GPL-3.0-or-later)'
####################################################

'''

 Usage:
	py -3 main.py
'''

import re
import json
from pprint import pprint
from collections import OrderedDict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

# global vars
data = OrderedDict()
doc = None
heading1_style = None
heading2_style = None
normal_style = None
normal_text = None
table = None
row_styles = []
row_vars = {}

def load_template(filename):
    ''' Load the docx template. '''
    global doc, heading1_style, heading2_style, normal_style, normal_text, \
            row_styles, row_vars, table
    doc = Document(filename)
    heading1_style = doc.paragraphs[0].style
    normal_style = doc.paragraphs[1].style
    normal_text = doc.paragraphs[1].text
    heading2_style = doc.paragraphs[2].style 
    table = doc.tables[0]

    for i in range(4):
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    table._element.getparent().remove(table._element)
    print('--Loaded table template--')
    for row in range(len(table.rows)):
        left, right = [table.cell(row, col) for col in (0, 1)]
        row_styles.append(left.paragraphs[0].runs[0].style)
        row_vars[right.text] = left.text
        print(left.text + ' : ' + right.text)
    print('--')

def load_excel_data(xlsx_filename, json_filename=None):
    ''' Load data from the first sheet in the xlsx file, and (optionally)
        export it in JSON format. '''
    global data
    wb = load_workbook(xlsx_filename)
    ws = wb.active

    header = [str(cell.value) for cell in list(ws.iter_rows())[0]]
    try:
        color_idx = header.index('Color')
    except ValueError:
        print("ERROR: Cannot find column 'Color' in sheet. Aborting.")
        return
    colors = [str(cell.value) for cell in list(ws.iter_cols())[color_idx] if cell.value][1:]
    colors = list(OrderedDict.fromkeys(colors))
    print(colors)
    for color in colors:
        if not color in data:
            data[color] = []
        for row in list(ws.iter_rows())[1:]:
            row = [str(cell.value) if cell.value else '' for cell in row]
            if row[color_idx] == color:
                data[color].append(dict(zip(header[:color_idx] + header[color_idx+1:], row[:color_idx] + row[color_idx+1:])))
            
    if json_filename:
        with open(json_filename, 'w') as f:
            f.write(json.dumps(data, indent=4))

def merge_data_and_save(filename):
    ''' Create a docx file based on the template and insert data from
        the loaded sheet. Save the result to `filename`.'''
    for color in data.keys():
        doc.add_heading(color, level=1).style = heading1_style
        key = re.findall('«.+»', normal_text)
        para = normal_text
        if key:
            para = normal_text.replace(key[0], color.lower())
        p = doc.add_paragraph(text=para, style=normal_style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for fruit in data[color]:
            doc.add_heading(fruit['Fruit'], level=2).style = heading2_style
            new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns), style=table.style)
            for row, row_org in zip(new_table.rows, table.rows):
                total_width = sum([row.cells[i].width for i in range(2)])
                ratio = row_org.cells[0].width / (row_org.cells[0].width+row_org.cells[1].width)
                row.cells[0].width = total_width * ratio
                row.cells[1].width = total_width * (1-ratio)
            for row in range(len(table.rows)):
                var = table.cell(row, 1).text
                left, right = [new_table.cell(row, col) for col in (0, 1)]
                left.paragraphs[0].add_run(row_vars[var]).bold = True
                try:
                    key = re.findall('«(.+)»', var)[0]
                    right.paragraphs[0].text = table.cell(row, 1).text.replace(f'«{key}»', fruit[key])
                except IndexError:
                    right.paragraphs[0].text = table.cell(row, 1).text
                #print(left, right)
            doc.add_paragraph()				
        doc.add_page_break()
    doc.save(filename)

if __name__ == '__main__':    
    load_template('template - edited.docx')
    load_excel_data('Fruit.xlsx')
    merge_data_and_save('demo.docx')
